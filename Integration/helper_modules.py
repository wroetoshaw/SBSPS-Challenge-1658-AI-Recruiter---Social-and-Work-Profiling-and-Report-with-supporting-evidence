import re
import docx
import time
import spacy
import string
import calendar
import datefinder
import numpy as np
import pandas as pd
from github import Github
from random import shuffle
from itertools import chain
from collections import Counter
import plotly.graph_objects as go
import matplotlib.pyplot as plt
from wordcloud import WordCloud
from collections import defaultdict
from pyresparser import ResumeParser
from nltk.corpus import stopwords
import matplotlib.patches as mpatches
from pyresparser.utils import extract_text_from_docx
from ibm_watson import NaturalLanguageUnderstandingV1
from ibm_cloud_sdk_core.authenticators import IAMAuthenticator
from ibm_watson.natural_language_understanding_v1 import Features, KeywordsOptions
import matplotlib
matplotlib.use('Agg')

import warnings
warnings.filterwarnings("ignore")

unique_id = None
file_path = None
username = None
to_search_username = None
password = None


def global_init(file_path=None, username=None, to_search_username=None, password=None, unique_id=None):
    globals()['file_path'] = file_path
    globals()['username'] = username
    globals()['to_search_username'] = to_search_username
    globals()['password'] = password
    globals()['unique_id'] = unique_id


def find_seperation(l):
    """
    Find seperation between various sections
    """
    for i in range(0, len(l)-1):
        if l[i] == l[i+1] == '':
            return i
        
        
def pad_list(a, b):
    mx = max(len(a), len(b))
    len_a, len_b = len(a), len(b)
    
    if len_a < mx:
        a = [a]
        a.append([np.nan] * (mx-len_a))
        return list(chain(*a)), b
    
    b = [b]
    b.append([np.nan] * (mx-len_b))
    return a, list(chain(*b))

                
def get_section_wise(section="PROJECTS", file_path=None):
    """
    Get the text of the defined section along with the Analysed result.
    """
    # Read file
    doc = docx.Document(file_path)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    read_data_docx = fullText

    # Filter Project related lines
    projects = list()
    flag = 0
    for i, line in enumerate(read_data_docx):
        if section in line:
            projects = read_data_docx[i:]
            break

    # Get projects related text
    project_content = projects[:find_seperation(projects)]

    # Save the projects related text
    name = file_path.split('.')[0]
    ext = file_path.split('.')[1]
    document = docx.Document()
    p = document.add_paragraph('\n'.join(project_content))
    document.save(unique_id + '_' +str(section).lower() + '.'+ ext)

    # Parse again on the projects to find skills related to Projects
    data_projects = ResumeParser(unique_id + '_' +str(section).lower() + '.'+ ext).get_extracted_data()
    
    return project_content, data_projects


def get_age_group(content):
    matches = datefinder.find_dates(content)
    dates_exp = list()
    for match in matches:
        dates_exp.append(match)
    print(dates_exp)
    year_of_completion = int(max(dates_exp).year)
    birth_year = year_of_completion-21
    age = 2020-birth_year
    print(year_of_completion, birth_year, age)
    if age < 0:
        return "14-18"
    age_groups = {
        "14-19": list(range(14, 19)),
        "19-24": list(range(19, 24)),
        "24-29": list(range(24, 29)),
        "29-34": list(range(29, 34)),
        "34-39": list(range(34, 39)),
        "39-44": list(range(39, 44)),
        "44-49": list(range(44, 49)),
        "49-54": list(range(49, 54)),
        "54-59": list(range(54, 59)),
        "59-100": list(range(60, 100))
    }
    for key in age_groups.keys():
        if age in age_groups[key]:
            return key
            
    
def cal_experience(content):
    """
    Finds dates in text, and calculates work experiece
    """
    # Find dates
    matches = datefinder.find_dates('\n'.join(content))
    dates_exp = list()
    for match in matches:
        dates_exp.append(match)

    # Calculate experience in months
    exp_in_months = list()
    for i in range(0, len(dates_exp)-1, 2):
        d1 = dates_exp[i]
        d2 = dates_exp[i+1]
        exp_in_months.append(abs(int((d2.year - d1.year + (d2.month - d1.month)/12)*12)))
        
    return exp_in_months

def get_keywords_nlu(docs):
    """
    Get keywords for a given set of documents in a list
    """
    t1 = time.time()
    
    # authenticator = IAMAuthenticator('yCpIapGBfywvAulhvvODJ0eHzu0cKhxaZRbjyjMpHN6P')
    # natural_language_understanding = NaturalLanguageUnderstandingV1(
    #     version='2019-07-12',
    #     authenticator=authenticator
    # )
    # natural_language_understanding.set_service_url('https://gateway-lon.watsonplatform.net/natural-language-understanding/api')
    authenticator = IAMAuthenticator('8zxWBBpd86s5VFvhs6g01-rXxOogAcZssn6AyTaJTRi9')
    natural_language_understanding = NaturalLanguageUnderstandingV1(
        version='2019-07-12',
        authenticator=authenticator
    )
    natural_language_understanding.set_service_url('https://api.eu-gb.natural-language-understanding.watson.cloud.ibm.com/instances/0cb1fe8d-a9e0-4659-8f76-009c73ee91cc')
    print("New Credentials")
    
    ls = list()
    for loc, each_article in enumerate(docs):
        response = natural_language_understanding.analyze(
        text=each_article,
        features=Features(keywords=KeywordsOptions(sentiment=False,emotion=False,limit=15))).get_result()
        ls.append(response)
        time.sleep(1)

    
    tuple_list = list()
    for each_record in ls:
        temp = list()
        for each_set in each_record['keywords']:
            temp.append((each_set['text'], each_set['relevance']))
        tuple_list.append(temp)
    
    keywords = list()
    conf = list()
    for k in tuple_list:
        temp = list()
        temp2 = list()
        for i, j in k:
            temp.append(i)
            temp2.append(j)
        keywords.append(temp)
        conf.append(temp2)
        
    print("Time elapsed for Keyword Extraction:", time.time()-t1)
    
    return keywords, conf


def get_donut(value, unique_id, name):
    plt.figure(figsize=(3, 3))
    print("matplotlib", value)
    value = min(10, int(round(value*10, 0))) * 10
    size_of_groups=[value, 100-value]
    print(size_of_groups)
    # Create a pieplot
    plt.pie(size_of_groups, colors=[(99/255, 110/255, 250/255),'red'], labels=[str(value), str(100-value)])
    #plt.show()
    
    # add a circle at the center
    my_circle=plt.Circle( (0,0), 0.7, color='white')
    p=plt.gcf()
    p.gca().add_artist(my_circle)
    plt.savefig('static/' + 'ere' + unique_id + name + '.png')
    return 


def generate_skills_text(a, b):
    a, b = set(map(lambda x:x.lower(), a)), set(map(lambda x:x.lower(), b))
    project_skills = a.intersection(b)
    overall_skills = a.difference(b)

    skill_freq_count = dict()
    resume_text = extract_text_from_docx(file_path).lower()
    for skill in project_skills:
        skill_freq_count[skill.lower()] = resume_text.count(skill.lower())

    for skill in overall_skills:
        skill_freq_count[skill.lower()] = resume_text.count(skill.lower())

    keys = list(skill_freq_count.keys())
    for key in keys:
        if len(key) < 2:
            del skill_freq_count[key]

    skill_text = ""
    for k, v in skill_freq_count.items():
        skill_text += ' '+ ' '.join([k]*v)
    
    skills = skill_text.split()    
    shuffle(skills)
    skill_text = ' '.join(skills)
    
    return skill_text, overall_skills, project_skills


def find_color(word, overall_skills, project_skills):
    if word.lower() in overall_skills:
        return 'rgb(255, 255, 255)'
    return 'rgb(255, 54, 11)'


def wordcloud_generation(skill_text, overall_skills, project_skills, unique_id, name):
    wc = WordCloud(stopwords = None, max_words = 200, max_font_size = 100)

    wc.generate(skill_text.strip())


    wordcloud_layout = list()
    for (word, freq), fontsize, position, orientation, color in wc.layout_:
        color = find_color(word, overall_skills, project_skills)
        wordcloud_layout.append(((word, freq), fontsize, position, orientation, color))
    wc.layout_ = wordcloud_layout

    plt.figure(figsize=(10, 10))
    plt.imshow(wc, interpolation='bilinear')
    plt.axis("off")
    plt.title("Skills - (Skills in Project vs Overall Skills mentioned in Resume)\nRed - Skills in project, White - Overall Skills")
    plt.tight_layout()
    plt.savefig('static/' + 'here' + unique_id + name + '.png', dpi=300)
    return 
    

def pie_chart(labels, values):
    fig = go.Figure(data=[go.Pie(labels=labels, values=values)])
    fig.update_layout(
        title_text="Work Experience vs Project Experience <br>(In months)", 
        title_x=0.5,
        autosize=False,
        width=500,
        height=500,
        margin=dict(
            l=50,
            r=50,
            b=100,
            t=100,
            pad=4
        ),
    )
    
    return fig
    
    
def get_line_color(percentages):
    if np.mean(np.array(list(map(int, percentages[1:])))) < 70:
        return 'rgb(199, 0, 57 )', 'Non-ideal'
    if np.mean(np.array(list(map(int, percentages[1:])))) < 85:
        return 'rgba(231,107,243)', 'Fair'
    return 'rgba(0,100,80)', 'Ideal'


def plot_lines(x, percentages):
    x_rev = x[::-1]

    y1 = [0]+percentages
    y1_upper = list(np.array(y1) + 10)
    y1_lower = list(np.array(y1) - 10)
    y1_lower = y1_lower[::-1]

    fig = go.Figure()
    lc, ln = get_line_color(y1)

    fig.add_trace(go.Scatter(
        x=x+x_rev,
        y=y1_upper+y1_lower,
        fill='toself',
        fillcolor=lc[:-1]+',0.3)',
        line_color='rgba(255,255,255,0)',
        showlegend=True,
        name=ln,
    ))

    fig.add_trace(go.Scatter(
        x=x, y=y1,
        line_color="blue",
        name='Performance',
    ))


    fig.update_traces(mode='lines+markers')
    
    fig.update_layout(
        title_text="Academics Performance", 
        title_x=0.5,
        autosize=False,
        width=500,
        height=500,
        margin=dict(
            l=50,
            r=50,
            b=100,
            t=100,
            pad=4
        )
    )
    
    return fig

########### Validations ###########


def get_repos(username):
    g = Github("18023d8111663e12dbc02df8edd2688fecac57a4")
    user = g.get_user(username)

    names = list()
    desc = list()
    for repo in user.get_repos():
        names.append(repo.name)
        desc.append(repo.description)
    return names, desc


def preprocess_title(names):
    stopwords_custom = stopwords.words('english')
    stopwords_custom += list(calendar.month_name)
    stopwords_custom += list(calendar.month_abbr)
    stopwords_custom += ['\n', '\t', '\r']
    stopwords_custom += ['title', 'name', 'description', 'project']
    stopwords_custom = list(map(lambda x:x.lower(), stopwords_custom))
    final_names = list()
    for name in names:
        temp = list()
        for word in name.lower().split():
            word = word.translate(str.maketrans('', '', string.punctuation+'–'))
            word = ''.join([i for i in word if not i.isdigit()])
            if word not in stopwords_custom:
                temp.append(word)
        final_names.append(' '.join(temp))
    
    return final_names


def similarity(a, b, nlp):
    return nlp(a).similarity(nlp(b))


def jaccard_similarity(list1, list2):
    intersection = len(list(set(list1).intersection(list2)))
    union = (len(list1) + len(list2)) - intersection
    return float(intersection) / union


def match_titles(names, titles, desc, username, threshold = 0.75):
    p_names = list(map(lambda x:x.lower().replace('_', ' ').replace('-', ''), names))
    titles = list(map(lambda x:x.lower().replace('_', ' ').replace('-', ''), titles))
    
    nlp = spacy.load('en_core_web_sm')
    matches = list()
    for title in titles:
        for name, orig in zip(p_names, names):
            if similarity(title, name, nlp) > threshold:
                matches.append("https://www.github.com/" + username + '/' + orig)
                break
        else:
            for des, orig in zip(desc, names):
                if jaccard_similarity(title.lower().split(), name.lower().split()) > (threshold-0.30):
                    matches.append("https://www.github.com/" + username + '/' + orig)
            else:
                matches.append(None)
    return matches


def plot_pie_validation(a, b):
    fig = go.Figure(data=[go.Pie(labels=['Found Supporting Evidence', 'No Evidence'], values=[a, b])])
    fig.update_layout(
        title_text="Supporting Evidence vs No Supporting Evidence <br>(For Projects)", 
        title_x=0.5,
        autosize=False,
        width=500,
        height=500,
        margin=dict(
            l=50,
            r=50,
            b=100,
            t=100,
            pad=4
        ),
    )
    return fig

########### Work Profile ###########


import json
import requests
import numpy as np
import pandas as pd

import requests
import plotly.graph_objects as go
from requests.auth import HTTPBasicAuth



def get_github_data(username, password, to_search_username):
    authentication = HTTPBasicAuth(username, password)

    data = requests.get('https://api.github.com/users/' + to_search_username, auth = authentication)
    data = data.json()

    url = data['repos_url']
    page_no = 1
    repos_data = []
    while (True):
        response = requests.get(url, auth = authentication)
        response = response.json()
        repos_data = repos_data + response
        repos_fetched = len(response)
        print("Total repositories fetched: {}".format(repos_fetched))
        if (repos_fetched == 30):
            page_no = page_no + 1
            url = (data['repos_url'] + '?page=' + str(page_no)).encode("UTF-8")
        else:
            break

    repos_information = []
    for i, repo in enumerate(repos_data):
        if repo['fork']:
            continue
        data = []
        data.append(repo['id'])
        data.append(repo['name'])
        data.append(repo['description'])
        data.append(repo['created_at'])
        data.append(repo['updated_at'])
        data.append(repo['owner']['login'])
        data.append(repo['license']['name'] if repo['license'] != None else None)
        data.append(repo['has_wiki'])
        data.append(repo['forks_count'])
        data.append(repo['open_issues_count'])
        data.append(repo['stargazers_count'])
        data.append(repo['watchers_count'])
        data.append(repo['url'])
        data.append(repo['commits_url'].split("{")[0])
        data.append(repo['url'] + '/languages')
        repos_information.append(data)

    repos_df = pd.DataFrame(repos_information, columns = ['Id', 'Name', 'Description', 'Created on', 'Updated on', 
                                                          'Owner', 'License', 'Includes wiki', 'Forks count', 
                                                          'Issues count', 'Stars count', 'Watchers count',
                                                          'Repo URL', 'Commits URL', 'Languages URL'])
    print("Total Source Repos:", len(repos_df))
    print("Getting Languages!")
    for i in range(repos_df.shape[0]):
        response = requests.get(repos_df.loc[i, 'Languages URL'], auth = authentication)
        response = response.json()
        if response != {}:
            languages = []
            for key, value in response.items():
                languages.append(key)
            languages = ', '.join(languages)
            repos_df.loc[i, 'Languages'] = languages
        else:
            repos_df.loc[i, 'Languages'] = ""

    print("Getting URLs!")
    response = requests.get(repos_df.loc[0, 'Commits URL'], auth = authentication)
    print("Getting COMMITS!")
    commits_information = []
    for i in range(repos_df.shape[0]):
        url = repos_df.loc[i, 'Commits URL']
        page_no = 1
        while (True):
            response = requests.get(url, auth = authentication)
            response = response.json()
            for commit in response:
                commit_data = []
                if type(commit) == type('str') or type(commit) == 1.0:
                    continue
                commit_data.append(repos_df.loc[i, 'Id'])
                commit_data.append(commit['sha'])
                commit_data.append(commit['commit']['committer']['date'])
                commit_data.append(commit['commit']['message'])
                commits_information.append(commit_data)
            if (len(response) == 30):
                page_no = page_no + 1
                url = repos_df.loc[i, 'Commits URL'] + '?page=' + str(page_no)
            else:
                break

    commits_df = pd.DataFrame(commits_information, columns = ['Repo Id', 'Commit Id', 'Date', 'Message'])

    commits_df.to_csv(unique_id + '_' + 'commits_info.csv', index = False)
    repos_df.to_csv(unique_id + '_' + 'repos_info.csv', index = False)
    print("Extracted and Saved GitHub Data!")
    
    return len(commits_df)


def get_dataframes(to_search_username):
    repos = pd.read_csv(unique_id + '_' + 'repos_info.csv')
    commits = pd.read_csv(unique_id + '_' + 'commits_info.csv')

    commits_count = pd.DataFrame(pd.merge(repos, 
             commits, 
             left_on='Id', 
             right_on='Repo Id', 
             how = 'left').groupby('Id').size().reset_index())
    commits_count.columns = ['Id', 'Commits count']

    repos = pd.merge(repos, commits_count, on = 'Id')
    
    return repos, commits


def plotly_barchart(x, y, title=None, xaxis_title=None, yaxis_title=None, tickvals=None, ticktext=None, ytickvals=None, width=None):
    fig = go.Figure([go.Bar(x=x, y=y, width=width)])
    fig.update_layout(
            title_text=title, 
            title_x=0.5,
            xaxis_title=xaxis_title,
            yaxis_title=yaxis_title,
        )
    fig.update_xaxes(
        ticktext=ticktext,
        tickvals=tickvals,
    )
    fig.update_yaxes(
        tickvals=ytickvals,
    )
    fig.update_layout(
        autosize=False,
        width=500,
        height=500,
        margin=dict(
            l=50,
            r=50,
            b=100,
            t=100,
            pad=4
        )
    )
    return fig
